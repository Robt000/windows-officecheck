import os
import sys
import time
import socket
import json
import re
import tkinter as tk
from tkinter import ttk, messagebox, simpledialog
import threading
import uuid
import platform
import getpass
import docx
import PyPDF2
import openpyxl
import requests
from datetime import datetime
import base64
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC

import socket
import requests
import time

class DocumentScanner:
    def __init__(self, server_url, employee_name, department):
        self.server_url = server_url
        self.employee_name = employee_name
        self.department = department
        self.hostname = socket.gethostname()
        self.ip_address = socket.gethostbyname(self.hostname)
        self.machine_id = str(uuid.uuid4())
        self.sensitive_keywords = [
            "涉密", "机密", "绝密", "敏感", "内部", "保密", 
            "产品设计", "设计文档", "产品规格", "技术方案",
            "客户信息", "用户数据", "财务", "合同"
        ]
        self.supported_extensions = [
            '.txt', '.doc', '.docx', '.pdf', '.xls', '.xlsx', 
            '.ppt', '.pptx', '.csv', '.rtf', '.md'
        ]
        # 添加需要跳过的系统文件夹
        self.skip_directories = [
            "Windows", "Program Files", "Program Files (x86)", 
            "$Recycle.Bin", "System Volume Information",
            "ProgramData", "AppData", "MSOCache", "Recovery",
            "PerfLogs", "Boot", "Config.Msi", "Documents and Settings",
            "Intel", "Microsoft", "Temp", "tmp"
        ]
    
    def scan_directory(self, directory, callback=None):
        """扫描指定目录下的所有文档"""
        documents = []
        alerts = []
        
        for root, dirs, files in os.walk(directory):
            # 跳过系统文件夹
            dirs[:] = [d for d in dirs if d not in self.skip_directories]
            for file in files:
                file_path = os.path.join(root, file)
                extension = os.path.splitext(file)[1].lower()
                if extension in self.supported_extensions:
                    try:
                        # 提取文档元数据
                        metadata = self.extract_office_metadata(file_path, extension)
                        
                        # 提取文件内容
                        content = self.extract_content(file_path, extension)
                        
                        # 判断是否包含敏感关键词
                        has_sensitive_content = False
                        matched_keywords = []
                        for keyword in self.sensitive_keywords:
                            if keyword in content:
                                has_sensitive_content = True
                                matched_keywords.append(keyword)
                        
                        # 创建文档记录
                        doc_record = {
                            "path": file_path,
                            "file_path": file_path,
                            "name": os.path.basename(file_path),
                            "file_name": os.path.basename(file_path),
                            "size": os.path.getsize(file_path),
                            "file_size": os.path.getsize(file_path),
                            "extension": extension,
                            "file_type": extension,
                            "creation_time": datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S'),
                            "create_time": datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S'),
                            "modified_time": datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S'),
                            "modify_time": datetime.fromtimestamp(os.path.getmtime(file_path)).strftime('%Y-%m-%d %H:%M:%S'),
                            "owner": self.get_file_owner(file_path),
                            "has_sensitive_content": has_sensitive_content,
                            "matched_keywords": matched_keywords,
                            
                            # 添加元数据字段
                            "author": metadata.get("author", ""),
                            "last_saved_by": metadata.get("last_saved_by", ""),
                            "revision": metadata.get("revision", ""),
                            "version": metadata.get("version", ""),
                            "company": metadata.get("company", ""),
                            "manager": metadata.get("manager", ""),
                            "metadata": metadata  # 保留完整的元数据字典以备其他需要
                        }
                        
                        documents.append(doc_record)
                        
                        # 如果包含敏感内容，添加到告警列表
                        if has_sensitive_content:
                            alert_record = {
                                "path": file_path,
                                "file_path": file_path,
                                "name": os.path.basename(file_path),
                                "file_name": os.path.basename(file_path),
                                "matched_keywords": matched_keywords,
                                "severity": "高" if len(matched_keywords) > 2 else "中"
                            }
                            alerts.append(alert_record)
                        
                        # 回调函数用于更新进度
                        if callback:
                            callback(f"扫描: {file_path}")
                    except Exception as e:
                        if callback:
                            callback(f"处理文件出错: {file_path}, 错误: {str(e)}")
        
        return documents, alerts

    def scan_full_disk(self, callback=None):
        """全盘扫描，跳过系统文件夹"""
        drives = [f'{d}:\\' for d in 'ABCDEFGHIJKLMNOPQRSTUVWXYZ' if os.path.exists(f'{d}:\\')]
        for drive in drives:
            self.scan_directory(drive, callback)

    def extract_office_metadata(self, file_path, extension):
        """提取Office文档的元数据"""
        metadata = {}
        try:
            if extension in ['.doc', '.docx']:
                doc = docx.Document(file_path)
                core_properties = doc.core_properties
                metadata['author'] = core_properties.author or ""
                metadata['last_saved_by'] = core_properties.last_modified_by or ""
                metadata['revision'] = str(core_properties.revision) if core_properties.revision else ""
                metadata['version'] = str(core_properties.version) if core_properties.version else ""
                metadata['company'] = core_properties.company or ""
                metadata['manager'] = core_properties.manager or ""
                metadata['created'] = core_properties.created.strftime('%Y-%m-%d %H:%M:%S') if core_properties.created else ""
                metadata['last_printed'] = core_properties.last_printed.strftime('%Y-%m-%d %H:%M:%S') if core_properties.last_printed else ""
                metadata['modified'] = core_properties.modified.strftime('%Y-%m-%d %H:%M:%S') if core_properties.modified else ""
            elif extension in ['.xls', '.xlsx']:
                wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
                props = wb.properties
                metadata['author'] = props.creator or ""
                metadata['last_saved_by'] = props.lastModifiedBy or ""
                metadata['company'] = props.company or ""
                metadata['manager'] = "" # Excel没有manager属性
                metadata['revision'] = "" # Excel没有revision属性
                metadata['version'] = "" # Excel没有version属性
                metadata['created'] = props.created.strftime('%Y-%m-%d %H:%M:%S') if props.created else ""
                metadata['modified'] = props.modified.strftime('%Y-%m-%d %H:%M:%S') if props.modified else ""
            elif extension == '.pdf':
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    info = reader.metadata
                    if info:
                        metadata['author'] = info.author or ""
                        metadata['creator'] = info.creator or ""
                        metadata['producer'] = info.producer or ""
                        metadata['subject'] = info.subject or ""
                        metadata['title'] = info.title or ""
                        # PDF没有这些属性，设为空字符串
                        metadata['company'] = ""
                        metadata['manager'] = ""
                        metadata['revision'] = ""
                        metadata['version'] = ""
        except Exception as e:
            print(f"Error extracting metadata from {file_path}: {e}")
        
        # 确保所有基本元数据字段都存在
        required_fields = ['author', 'last_saved_by', 'company', 'manager', 'revision', 'version']
        for field in required_fields:
            if field not in metadata:
                metadata[field] = ""
                
        return metadata
    
    def get_file_owner(self, file_path):
        """获取文件所有者"""
        try:
            if platform.system() == 'Windows':
                import win32security
                sd = win32security.GetFileSecurity(file_path, win32security.OWNER_SECURITY_INFORMATION)
                owner_sid = sd.GetSecurityDescriptorOwner()
                name, domain, type = win32security.LookupAccountSid(None, owner_sid)
                return f"{domain}\\{name}"
            else:
                import pwd
                stat_info = os.stat(file_path)
                uid = stat_info.st_uid
                return pwd.getpwuid(uid).pw_name
        except:
            return getpass.getuser()  # 如果获取失败，返回当前用户
    
    def extract_content(self, file_path, extension):
        """提取文件内容"""
        try:
            if extension in ['.txt', '.md', '.csv', '.rtf']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            
            elif extension == '.pdf':
                content = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page_num in range(len(pdf_reader.pages)):
                        content += pdf_reader.pages[page_num].extract_text()
                return content
            
            elif extension == '.docx':
                doc = docx.Document(file_path)
                return " ".join([para.text for para in doc.paragraphs])
            
            elif extension == '.xlsx':
                workbook = openpyxl.load_workbook(file_path, read_only=True)
                content = ""
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    for row in sheet.iter_rows(values_only=True):
                        content += " ".join([str(cell) if cell is not None else "" for cell in row]) + " "
                return content
            
            # 其他格式暂不支持内容提取
            return ""
        except Exception as e:
            print(f"提取内容出错: {file_path}, 错误: {str(e)}")
            return ""
    
    def send_data_to_server(self, documents, alerts, callback=None, scan_start_time=None, scan_end_time=None):
        """发送数据到服务器"""
        try:
            # 自定义JSON编码器，处理datetime对象
            class DateTimeEncoder(json.JSONEncoder):
                def default(self, obj):
                    if isinstance(obj, datetime):
                        return obj.strftime('%Y-%m-%d %H:%M:%S')
                    return super().default(obj)
            
            # 递归处理文档中的字段名称，确保兼容性
            def process_documents(docs):
                for doc in docs:
                    # 重命名字段以确保兼容性
                    if 'create_time' in doc and 'creation_time' not in doc:
                        doc['creation_time'] = doc['create_time']
                        del doc['create_time']
                    if 'modify_time' in doc and 'modified_time' not in doc:
                        doc['modified_time'] = doc['modify_time']
                        del doc['modify_time']
                    if 'modification_time' in doc and 'modified_time' not in doc:
                        doc['modified_time'] = doc['modification_time']
                        del doc['modification_time']
                    if 'file_type' in doc and 'extension' not in doc:
                        doc['extension'] = doc['file_type']
                        del doc['file_type']
                    if 'file_size' in doc and 'size' not in doc:
                        doc['size'] = doc['file_size']
                        del doc['file_size']
                    if 'file_path' in doc and 'path' not in doc:
                        doc['path'] = doc['file_path']
                        del doc['file_path']
                    
                    # 确保同时有name和file_name两个字段
                    if 'name' in doc and 'file_name' not in doc:
                        doc['file_name'] = doc['name']
                    elif 'file_name' in doc and 'name' not in doc:
                        doc['name'] = doc['file_name']
                return docs
            
            # 处理告警字段名称
            def process_alerts(alert_list):
                for alert in alert_list:
                    if 'file_path' in alert and 'path' not in alert:
                        alert['path'] = alert['file_path']
                        del alert['file_path']
                    
                    # 确保同时有name和file_name两个字段
                    if 'name' in alert and 'file_name' not in alert:
                        alert['file_name'] = alert['name']
                    elif 'file_name' in alert and 'name' not in alert:
                        alert['name'] = alert['file_name']
                return alert_list
            
            # 获取当前时间作为detection_time
            detection_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # 计算扫描耗时（如果提供了开始和结束时间）
            scan_duration = None
            if scan_start_time and scan_end_time:
                scan_duration = (scan_end_time - scan_start_time).total_seconds()
            
            data = {
                "machine_id": self.machine_id,
                "hostname": self.hostname,
                "ip_address": self.ip_address,
                "employee_name": self.employee_name,
                "department": self.department,
                "scan_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "detection_time": detection_time,  # 添加detection_time字段
                "scan_start_time": scan_start_time.strftime('%Y-%m-%d %H:%M:%S') if scan_start_time else None,
                "scan_end_time": scan_end_time.strftime('%Y-%m-%d %H:%M:%S') if scan_end_time else None,
                "scan_duration": scan_duration,  # 扫描耗时（秒）
                "documents": process_documents(documents),
                "alerts": process_alerts(alerts)
            }
            
            # 使用自定义编码器进行JSON序列化
            json_data = json.dumps(data, cls=DateTimeEncoder)
            
            response = requests.post(
                f"{self.server_url}/api/submit", 
                data=json_data,
                headers={'Content-Type': 'application/json'}
            )
            
            if response.status_code == 200:
                if callback:
                    callback(f"数据已成功发送到服务器")
                return True
            else:
                if callback:
                    callback(f"发送数据失败: {response.text}")
                return False
        except Exception as e:
            if callback:
                callback(f"发送数据出错: {str(e)}")
            return False

    def test_server_connection(self, callback=None):
        """测试与服务器的连接"""
        if callback:
            callback("正在测试与服务器的连接...")
        
        try:
            # 解析服务器URL获取主机和端口
            from urllib.parse import urlparse
            parsed_url = urlparse(self.server_url)
            host = parsed_url.hostname
            port = parsed_url.port or 5000  # 默认端口为5000
            
            # 测试TCP连接
            s = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
            s.settimeout(5)  # 设置超时时间为5秒
            result = s.connect_ex((host, port))
            s.close()
            
            if result == 0:
                # TCP连接成功，再测试HTTP连接
                response = requests.get(f"{self.server_url}/api/ping", timeout=5)
                if response.status_code == 200:
                    if callback:
                        callback("服务器连接测试成功！")
                    return True
                else:
                    if callback:
                        callback(f"HTTP连接失败: 状态码 {response.status_code}")
                    return False
            else:
                if callback:
                    callback(f"TCP连接失败: 无法连接到 {host}:{port}")
                return False
                
        except Exception as e:
            if callback:
                callback(f"连接测试出错: {str(e)}")
            return False

    def export_scan_data(self, documents, alerts, password, export_path=None, callback=None, scan_start_time=None, scan_end_time=None):
        """导出扫描数据到加密文件"""
        try:
            if callback:
                callback("正在准备导出数据...")
            
            # 自定义JSON编码器，处理datetime对象
            class DateTimeEncoder(json.JSONEncoder):
                def default(self, obj):
                    if isinstance(obj, datetime):
                        return obj.strftime('%Y-%m-%d %H:%M:%S')
                    return super().default(obj)
            
            # 递归处理文档中的字段名称，确保兼容性
            def process_documents(docs):
                for doc in docs:
                    # 重命名字段以确保兼容性
                    if 'create_time' in doc and 'creation_time' not in doc:
                        doc['creation_time'] = doc['create_time']
                        del doc['create_time']
                    if 'modify_time' in doc and 'modified_time' not in doc:
                        doc['modified_time'] = doc['modify_time']
                        del doc['modify_time']
                    if 'modification_time' in doc and 'modified_time' not in doc:
                        doc['modified_time'] = doc['modification_time']
                        del doc['modification_time']
                    if 'file_type' in doc and 'extension' not in doc:
                        doc['extension'] = doc['file_type']
                        del doc['file_type']
                    if 'file_size' in doc and 'size' not in doc:
                        doc['size'] = doc['file_size']
                        del doc['file_size']
                    if 'file_path' in doc and 'path' not in doc:
                        doc['path'] = doc['file_path']
                        del doc['file_path']
                    
                    # 确保同时有name和file_name两个字段
                    if 'name' in doc and 'file_name' not in doc:
                        doc['file_name'] = doc['name']
                    elif 'file_name' in doc and 'name' not in doc:
                        doc['name'] = doc['file_name']
                return docs
            
            # 处理告警字段名称
            def process_alerts(alert_list):
                for alert in alert_list:
                    if 'file_path' in alert and 'path' not in alert:
                        alert['path'] = alert['file_path']
                        del alert['file_path']
                    
                    # 确保同时有name和file_name两个字段
                    if 'name' in alert and 'file_name' not in alert:
                        alert['file_name'] = alert['name']
                    elif 'file_name' in alert and 'name' not in alert:
                        alert['name'] = alert['file_name']
                return alert_list
            
            # 获取当前时间作为detection_time
            detection_time = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            
            # 计算扫描耗时（如果提供了开始和结束时间）
            scan_duration = None
            if scan_start_time and scan_end_time:
                scan_duration = (scan_end_time - scan_start_time).total_seconds()
            
            # 准备导出数据
            export_data = {
                "machine_id": self.machine_id,
                "hostname": self.hostname,
                "ip_address": self.ip_address,
                "employee_name": self.employee_name,
                "department": self.department,
                "scan_time": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                "detection_time": detection_time,  # 添加detection_time字段
                "scan_start_time": scan_start_time.strftime('%Y-%m-%d %H:%M:%S') if scan_start_time else None,
                "scan_end_time": scan_end_time.strftime('%Y-%m-%d %H:%M:%S') if scan_end_time else None,
                "scan_duration": scan_duration,  # 扫描耗时（秒）
                "documents": process_documents(documents),
                "alerts": process_alerts(alerts)
            }
            
            # 将数据转换为JSON，使用自定义编码器
            json_data = json.dumps(export_data, ensure_ascii=False, cls=DateTimeEncoder)
            
            # 生成加密密钥
            salt = os.urandom(16)
            kdf = PBKDF2HMAC(
                algorithm=hashes.SHA256(),
                length=32,
                salt=salt,
                iterations=100000,
            )
            key = base64.urlsafe_b64encode(kdf.derive(password.encode()))
            cipher = Fernet(key)
            
            # 加密数据
            encrypted_data = cipher.encrypt(json_data.encode())
            
            # 如果没有指定导出路径，则在当前目录创建
            if not export_path:
                export_path = os.path.join(os.getcwd(), f"scan_data_{datetime.now().strftime('%Y%m%d_%H%M%S')}.sdat")
            
            # 创建加密文件
            with open(export_path, 'wb') as f:
                # 写入salt
                f.write(salt)
                # 写入加密数据
                f.write(encrypted_data)
            
            if callback:
                callback(f"数据已成功导出到: {export_path}")
            
            return export_path
        
        except Exception as e:
            if callback:
                callback(f"导出数据失败: {str(e)}")
            return None


class ScannerGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("科来信息安全审计工具")
        self.root.geometry("600x500")
        self.root.resizable(True, True)
        
        # 服务器地址
        self.server_frame = ttk.LabelFrame(root, text="服务器配置")
        self.server_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(self.server_frame, text="服务器地址:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.server_url = tk.StringVar(value="http://192.168.110.2:5000")
        server_entry = ttk.Entry(self.server_frame, textvariable=self.server_url, width=40)
        server_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        server_entry.config(state="readonly")  # 设置为只读，防止修改
        
        # 用户信息
        self.user_frame = ttk.LabelFrame(root, text="用户信息")
        self.user_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(self.user_frame, text="员工姓名:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.employee_name = tk.StringVar()
        ttk.Entry(self.user_frame, textvariable=self.employee_name, width=20).grid(row=0, column=1, padx=5, pady=5, sticky="w")
        
        ttk.Label(self.user_frame, text="部门:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        self.department = tk.StringVar()
        self.departments = ["技术部", "市场部", "销售部", "人力资源部", "财务部", "行政部", "其他"]
        ttk.Combobox(self.user_frame, textvariable=self.department, values=self.departments, width=18).grid(row=1, column=1, padx=5, pady=5, sticky="w")
        
        # 扫描配置
        self.scan_frame = ttk.LabelFrame(root, text="扫描配置")
        self.scan_frame.pack(fill="x", padx=10, pady=5)
        
        ttk.Label(self.scan_frame, text="扫描路径:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        self.scan_path = tk.StringVar(value="C:/Users/admin/Desktop")
        scan_entry = ttk.Entry(self.scan_frame, textvariable=self.scan_path, width=40)
        scan_entry.grid(row=0, column=1, padx=5, pady=5, sticky="w")
        scan_entry.config(state="readonly")  # 设置为只读
        ttk.Button(self.scan_frame, text="浏览", command=self.browse_directory).grid(row=0, column=2, padx=5, pady=5)
        
        # 日志区域
        self.log_frame = ttk.LabelFrame(root, text="扫描日志")
        self.log_frame.pack(fill="both", expand=True, padx=10, pady=5)
        
        self.log_text = tk.Text(self.log_frame, wrap=tk.WORD, height=10)
        self.log_text.pack(fill="both", expand=True, padx=5, pady=5)
        
        scrollbar = ttk.Scrollbar(self.log_text, command=self.log_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.log_text.config(yscrollcommand=scrollbar.set)
        
        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress = ttk.Progressbar(root, variable=self.progress_var, maximum=100)
        self.progress.pack(fill="x", padx=10, pady=5)
        
        # 按钮区域
        self.button_frame = ttk.Frame(root)
        self.button_frame.pack(fill="x", padx=10, pady=10)
        
        self.scan_button = ttk.Button(self.button_frame, text="开始扫描", command=self.start_scan)
        self.scan_button.pack(side=tk.RIGHT, padx=5)
        
        self.export_button = ttk.Button(self.button_frame, text="导出数据", command=self.export_data, state=tk.DISABLED)
        self.export_button.pack(side=tk.RIGHT, padx=5)
        
        self.exit_button = ttk.Button(self.button_frame, text="退出", command=root.destroy)
        self.exit_button.pack(side=tk.RIGHT, padx=5)
        
        # 扫描线程
        self.scan_thread = None
        self.is_scanning = False
        
        # 初始化上次扫描结果
        self.last_scan_documents = []
        self.last_scan_alerts = []
    
    def browse_directory(self):
        """选择扫描目录"""
        from tkinter import filedialog
        directory = filedialog.askdirectory()
        if directory:
            self.scan_path.set(directory)
    
    def log(self, message):
        """添加日志"""
        self.log_text.insert(tk.END, f"{datetime.now().strftime('%H:%M:%S')} - {message}\n")
        self.log_text.see(tk.END)
    
    def start_scan(self):
        """开始扫描"""
        if self.is_scanning:
            messagebox.showinfo("提示", "扫描正在进行中，请等待完成")
            return
        
        employee_name = self.employee_name.get().strip()
        department = self.department.get().strip()
        server_url = self.server_url.get().strip()
        scan_path = self.scan_path.get()
        
        if not employee_name:
            messagebox.showerror("错误", "请输入员工姓名")
            return
        
        if not department:
            messagebox.showerror("错误", "请选择部门")
            return
        
        # 初始化扫描器
        scanner = DocumentScanner(server_url, employee_name, department)
        
        # 测试服务器连接
        self.log("正在测试与服务器的连接...")
        if not scanner.test_server_connection(self.log):
            if not messagebox.askyesno("警告", "无法连接到服务器，是否继续扫描？\n(扫描结果将无法发送到服务器)"):
                return
        
        self.is_scanning = True
        self.scan_button.config(state=tk.DISABLED)
        self.progress_var.set(0)
        self.log_text.delete(1.0, tk.END)
        
        # 更新日志显示实际扫描路径
        if scan_path == "所有盘符":
            self.log(f"开始扫描所有盘符")
        else:
            self.log(f"开始扫描路径: {scan_path}")
        
        self.scan_thread = threading.Thread(target=self.perform_scan, args=(server_url, employee_name, department))
        self.scan_thread.daemon = True
        self.scan_thread.start()
    
    def perform_scan(self, server_url, employee_name, department):
        """执行扫描操作"""
        try:
            scanner = DocumentScanner(server_url, employee_name, department)
            self.log(f"初始化扫描器完成")
            
            # 获取扫描路径
            scan_path = self.scan_path.get()
            
            # 记录扫描开始时间
            scan_start_time = datetime.now()
            self.log(f"扫描开始时间: {scan_start_time.strftime('%Y-%m-%d %H:%M:%S')}")
            
            if scan_path == "所有盘符":
                # 获取所有可用盘符
                available_drives = []
                for drive in range(ord('A'), ord('Z')+1):
                    drive_letter = chr(drive) + ":\\"
                    if os.path.exists(drive_letter):
                        available_drives.append(drive_letter)
                
                self.log(f"发现可用盘符: {', '.join(available_drives)}")
                
                # 计算所有盘符中的文件总数
                total_files = 0
                scanned_files = 0
                
                for drive in available_drives:
                    self.log(f"正在统计 {drive} 中的文档文件...")
                    try:
                        for root, _, files in os.walk(drive):
                            for file in files:
                                _, ext = os.path.splitext(file)
                                if ext.lower() in scanner.supported_extensions:
                                    total_files += 1
                    except Exception as e:
                        self.log(f"统计 {drive} 时出错: {str(e)}")
            else:
                # 指定目录扫描
                self.log(f"扫描路径: {scan_path}")
                
                # 计算指定目录中的文件总数
                total_files = 0
                scanned_files = 0
                
                try:
                    for root, _, files in os.walk(scan_path):
                        for file in files:
                            _, ext = os.path.splitext(file)
                            if ext.lower() in scanner.supported_extensions:
                                total_files += 1
                except Exception as e:
                    self.log(f"统计文件数量时出错: {str(e)}")
            
            if total_files == 0:
                self.log("未找到支持的文档文件")
                self.root.after(0, self.finish_scan)
                return
            
            self.log(f"找到 {total_files} 个文档文件待扫描")
            
            # 定义回调函数更新进度
            def update_progress(message):
                nonlocal scanned_files
                scanned_files += 1
                progress = (scanned_files / total_files) * 100
                self.root.after(0, lambda: self.progress_var.set(progress))
                self.root.after(0, lambda: self.log(message))
            
            # 执行扫描
            self.log("开始扫描文档...")
            documents = []
            alerts = []
            
            if scan_path == "所有盘符":
                # 扫描所有盘符
                for drive in available_drives:
                    self.log(f"正在扫描 {drive}...")
                    try:
                        drive_documents, drive_alerts = scanner.scan_directory(drive, update_progress)
                        documents.extend(drive_documents)
                        alerts.extend(drive_alerts)
                    except Exception as e:
                        self.log(f"扫描 {drive} 时出错: {str(e)}")
            else:
                # 扫描指定目录
                try:
                    dir_documents, dir_alerts = scanner.scan_directory(scan_path, update_progress)
                    documents.extend(dir_documents)
                    alerts.extend(dir_alerts)
                except Exception as e:
                    self.log(f"扫描 {scan_path} 时出错: {str(e)}")
            
            # 记录扫描结束时间
            scan_end_time = datetime.now()
            scan_duration = (scan_end_time - scan_start_time).total_seconds()
            self.log(f"扫描结束时间: {scan_end_time.strftime('%Y-%m-%d %H:%M:%S')}")
            self.log(f"扫描总耗时: {int(scan_duration//60)}分{int(scan_duration%60)}秒")
            
            # 保存扫描结果，用于导出功能
            self.last_scan_documents = documents
            self.last_scan_alerts = alerts
            self.last_scan_start_time = scan_start_time
            self.last_scan_end_time = scan_end_time
            
            # 发送数据到服务器
            self.log(f"扫描完成，共发现 {len(documents)} 个文档，{len(alerts)} 个告警")
            self.log("正在将数据发送到服务器...")
            
            success = scanner.send_data_to_server(documents, alerts, self.log, scan_start_time, scan_end_time)
            
            if success:
                self.log("数据已成功发送到服务器")
            else:
                self.log("发送数据到服务器失败")
            
            # 完成扫描
            self.root.after(0, self.finish_scan)
            
        except Exception as e:
            self.log(f"扫描过程出错: {str(e)}")
            self.root.after(0, self.finish_scan)
    
    def finish_scan(self):
        """完成扫描，重置界面状态"""
        self.is_scanning = False
        self.scan_button.config(state=tk.NORMAL)
        self.export_button.config(state=tk.NORMAL)  # 启用导出按钮
        self.log("扫描任务结束")

    def export_data(self):
        """导出扫描数据"""
        if not self.last_scan_documents and not self.last_scan_alerts:
            messagebox.showinfo("提示", "没有可导出的扫描数据")
            return
        
        # 使用默认密码，不再询问用户
        password = "Csxxb@2025"
        
        # 选择保存位置
        from tkinter import filedialog
        export_path = filedialog.asksaveasfilename(
            defaultextension=".sdat",
            filetypes=[("扫描数据文件", "*.sdat"), ("所有文件", "*.*")],
            title="保存扫描数据"
        )
        if not export_path:
            return
        
        # 初始化扫描器
        server_url = self.server_url.get().strip()
        employee_name = self.employee_name.get().strip()
        department = self.department.get().strip()
        scanner = DocumentScanner(server_url, employee_name, department)
        
        # 导出数据
        self.log("正在导出扫描数据...")
        
        # 获取扫描时间信息
        scan_start_time = getattr(self, 'last_scan_start_time', None)
        scan_end_time = getattr(self, 'last_scan_end_time', None)
        
        export_path = scanner.export_scan_data(
            self.last_scan_documents, 
            self.last_scan_alerts, 
            password, 
            export_path, 
            self.log,
            scan_start_time,
            scan_end_time
        )
        
        if export_path:
            messagebox.showinfo("成功", f"数据已成功导出到:\n{export_path}")


def main():
    root = tk.Tk()
    app = ScannerGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
